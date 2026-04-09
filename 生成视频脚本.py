from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def make_run(p, text, font_size=12, bold=False, color=None, font_name="宋体"):
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_para(doc, text="", font_size=12, bold=False, align=None, color=None, space_before=2, space_after=4, font_name="宋体"):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    make_run(p, text, font_size, bold, color, font_name)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)
    return p

def add_divider(doc):
    add_para(doc, "─" * 40, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(180, 180, 180))

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ===== 封面标题 =====
add_para(doc, "", font_size=6)
add_para(doc, "华香古方 · 香文化短视频脚本", font_size=22, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(102, 51, 0))
add_para(doc, "— 适用于剪映「图文成片」/「AI成片」 —", font_size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(150, 120, 80))
add_para(doc, "", font_size=6)

add_divider(doc)

# ===== 使用说明 =====
add_para(doc, "使用方法", font_size=14, bold=True, color=(102, 51, 0))
steps = [
    "1. 打开「剪映」App → 点击首页「图文成片」或「AI成片」",
    "2. 将下方【旁白文案】整段复制粘贴进去",
    "3. 剪映会自动匹配画面、生成字幕和配音",
    "4. 生成后可微调：替换为自己实拍的素材效果更佳",
    "5. 推荐选择「古风」「国潮」「禅意」风格模板",
    "6. 配乐建议：古琴、箫、空灵鼓等中式纯音乐",
]
for s in steps:
    add_para(doc, s, font_size=11, space_after=2)

add_divider(doc)

# ===== 脚本一 =====
add_para(doc, "脚本一：《暗香千年 · 古法线香的秘密》", font_size=15, bold=True, color=(102, 51, 0), space_before=10)
add_para(doc, "时长：约 45-60 秒  |  风格：禅意治愈  |  标签：#香文化 #线香 #非遗 #东方美学", font_size=10, color=(120, 120, 120))

add_para(doc, "【旁白文案 — 复制以下内容到剪映】", font_size=12, bold=True, space_before=8)

script1 = """你知道吗，中国人燃香的历史，已经超过三千年了。

一缕青烟，从指尖升起，穿过千年时光，连接着我们与古人的对话。

古法制香，不用一滴化学香精。
沉香、檀香、降真香、中草药，每一味香材，都来自大自然的馈赠。

经过选材、研磨、调配、搓制、晾晒，
一根线香的诞生，需要经历七十二道工序。

点燃它，你会看到——
烟雾缓缓升起，凝聚成云，这就是古人所说的"香云盖"。

在这个快节奏的时代，
不妨点一炷香，让自己慢下来。

暗香一抹，氤氲千年。
与香为友，让身心清净。"""

add_para(doc, script1, font_size=12, space_before=4, space_after=6)

add_para(doc, "【画面建议 — 供参考，剪映会自动匹配】", font_size=12, bold=True)
scenes1 = [
    "画面1：古寺晨光 / 青铜香炉特写 → 配文字「三千年的传承」",
    "画面2：青烟从香炉缓缓升起的慢镜头",
    "画面3：各种天然香料摆放（沉香木、檀香粉、草药）",
    "画面4：手工搓制线香的过程特写",
    "画面5：线香整齐排列晾晒的画面",
    "画面6：点燃线香，烟雾上升形成香云盖的特写",
    "画面7：一个人在窗前静坐品香的背影，阳光洒入",
    "画面8：结尾字幕「暗香一抹 氤氲千年」+ 账号logo",
]
for s in scenes1:
    add_para(doc, s, font_size=10, color=(80, 80, 80), space_after=2)

add_divider(doc)

# ===== 脚本二 =====
add_para(doc, "脚本二：《一炷香的时间》", font_size=15, bold=True, color=(102, 51, 0), space_before=10)
add_para(doc, "时长：约 30-45 秒  |  风格：治愈清新  |  标签：#香文化 #手工香 #禅意生活 #慢生活", font_size=10, color=(120, 120, 120))

add_para(doc, "【旁白文案 — 复制以下内容到剪映】", font_size=12, bold=True, space_before=8)

script2 = """古人计时，不用钟表，用一炷香。

一炷香燃尽，大约是半个时辰。
在这段时间里，他们读书、抚琴、煮茶、静思。

你有多久，没有给自己这样一段安静的时光了？

选一款天然古方线香，
看烟雾旋转、升腾、消散，
把心里的浮躁，一点一点放下。

不必远行，不必花费太多，
一炷香的时间，就能让你回归内心的宁静。

给自己一炷香的时间，
去感受生活本来的样子。"""

add_para(doc, script2, font_size=12, space_before=4, space_after=6)

add_para(doc, "【画面建议 — 供参考，剪映会自动匹配】", font_size=12, bold=True)
scenes2 = [
    "画面1：古代沙漏或日晷 → 转场到一炷燃烧中的线香",
    "画面2：古风场景：书桌、毛笔、茶杯旁的香炉",
    "画面3：现代人在忙碌城市中快步行走（对比）",
    "画面4：回到室内，手轻轻点燃一根线香",
    "画面5：烟雾螺旋上升的慢镜头特写",
    "画面6：一杯茶旁边放着香炉，窗外是绿植",
    "画面7：人闭眼微笑，享受宁静时刻",
    "画面8：结尾字幕「一炷香的时间 回归宁静」+ 账号logo",
]
for s in scenes2:
    add_para(doc, s, font_size=10, color=(80, 80, 80), space_after=2)

add_divider(doc)

# ===== 脚本三 =====
add_para(doc, "脚本三：《你烧的香，是天然的吗？》", font_size=15, bold=True, color=(102, 51, 0), space_before=10)
add_para(doc, "时长：约 40-50 秒  |  风格：科普种草  |  标签：#香文化 #天然线香 #华香古方香文化 #传统文化", font_size=10, color=(120, 120, 120))

add_para(doc, "【旁白文案 — 复制以下内容到剪映】", font_size=12, bold=True, space_before=8)

script3 = """很多人问我，天然香和化学香，到底有什么区别？

其实很简单。

化学香，点燃后气味刺鼻，烟雾发黑，燃烧后香灰散乱。
闻久了，会头晕、不舒服。

而天然古方线香，用的是沉香、檀香、草本植物。
点燃后，香气清雅悠长，烟色洁白。
香灰不散不落，甚至能看到传说中的"香云盖"。

好的香，不是掩盖气味，而是安抚你的心神。

我们家族传承古法制香一百六十多年，
每一根线香，都是手工搓制，自然晾干。
只为守住这份来自大自然的纯粹。

闻过真正的好香，你就再也回不去了。"""

add_para(doc, script3, font_size=12, space_before=4, space_after=6)

add_para(doc, "【画面建议 — 供参考，剪映会自动匹配】", font_size=12, bold=True)
scenes3 = [
    "画面1：两种香对比摆放（左：颜色鲜艳的廉价香 vs 右：素雅天然线香）",
    "画面2：化学香燃烧 → 黑烟、刺鼻（画面可用红色X标记）",
    "画面3：天然线香燃烧 → 白烟袅袅、香灰挺立",
    "画面4：各种天然香材特写（沉香、檀香原木、草药）",
    "画面5：手工搓制线香的过程",
    "画面6：香云盖形成的震撼画面",
    "画面7：老师傅制香 / 传承场景",
    "画面8：结尾「百年传承 只为一缕真香」+ 账号logo + 购买引导",
]
for s in scenes3:
    add_para(doc, s, font_size=10, color=(80, 80, 80), space_after=2)

add_divider(doc)

# ===== 发布建议 =====
add_para(doc, "发布建议", font_size=14, bold=True, color=(102, 51, 0), space_before=10)
tips = [
    "发布时间：建议早上 7-9 点 或 晚上 8-10 点（流量高峰）",
    "封面：选烟雾特写或香云盖画面，加上大字标题",
    "配乐：剪映搜索「古风纯音乐」「禅意」「空灵鼓」",
    "话题标签：#香文化 #线香 #手工香 #东方美学 #非遗 #传统文化 #禅意生活 #治愈",
    "互动引导：结尾加「你喜欢什么香？评论区告诉我」提升评论量",
    "合集建议：创建「每日香文化分享」合集，持续更新形成系列",
]
for t in tips:
    add_para(doc, "· " + t, font_size=11, space_after=3)

output_path = "/mnt/d/Claude Code Data/华香古方_短视频脚本_剪映图文成片.docx"
doc.save(output_path)
print(f"文档已保存: {output_path}")
