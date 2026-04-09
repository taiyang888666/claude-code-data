from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def make_run(p, text, font_size=12, bold=False, color=None, font_name="宋体"):
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_para(doc, text="", font_size=12, bold=False, align=None, color=None, space_before=2, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    make_run(p, text, font_size, bold, color)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)
    return p

def add_divider(doc):
    add_para(doc, "─" * 45, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(180, 180, 180))

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# 标题
add_para(doc, "无穷盐焗鸡腿 · 抖音带货短视频脚本", font_size=20, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(200, 50, 0))
add_para(doc, "即梦AI生成 + 剪映拼接 | 抖音合规版", font_size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(150, 120, 80))

add_divider(doc)

# ===== 全局设定 =====
add_para(doc, "一、全局设定", font_size=15, bold=True, color=(200, 50, 0), space_before=8)

settings = [
    ("画面比例", "9:16 竖屏短视频"),
    ("视觉风格", "iPhone后置原相机实拍感，自然温暖生活化打光，高清晰度，真实质感"),
    ("剪辑节奏", "18秒内高频跳切，BGM为动感卡点打击乐"),
    ("音效", "撕包装声、大口咀嚼声、吞咽音效（清晰自然）"),
    ("人物设定", "二十多岁年轻吃货，纯色卫衣，自然皮肤质感，吃相有感染力"),
    ("总时长", "约15-18秒"),
]
for label, value in settings:
    p = doc.add_paragraph()
    make_run(p, f"  {label}：", 11, bold=True)
    make_run(p, value, 11)
    p.paragraph_format.space_after = Pt(2)

add_divider(doc)

# ===== 抖音合规提醒 =====
add_para(doc, "二、抖音平台合规要点", font_size=15, bold=True, color=(200, 50, 0), space_before=8)

rules = [
    "1. 不使用绝对化用语（最好吃、第一、唯一等）",
    "2. 不涉及虚假功效宣传（如减肥、养生、治病等）",
    "3. 食品类广告需标注广告标识，如带货需挂车",
    "4. 不出现竞品贬低或对比",
    "5. 口播文案不含低俗、暴力、色情内容",
    "6. 花字/字幕中不出现违禁词（可用剪映自带的违禁词检测）",
    "7. 背景音乐需使用剪映曲库中有版权的音乐",
    "8. 画面不能过度夸张变形，保持真实感",
]
for r in rules:
    add_para(doc, r, font_size=11, space_after=2)

add_divider(doc)

# ===== 分镜脚本 =====
add_para(doc, "三、分镜脚本（合规版）", font_size=15, bold=True, color=(200, 50, 0), space_before=8)

# 镜头1
add_para(doc, "镜头1：视觉暴击与开箱（00:00 - 00:03）", font_size=13, bold=True, color=(50, 50, 50), space_before=10)

add_para(doc, "【口播文案】（急促激动，自然语气）", font_size=11, bold=True)
add_para(doc, '"饿了馋了？必须来个无穷盐焗大鸡腿！"', font_size=12, color=(200, 50, 0))

add_para(doc, "【即梦AI提示词 — 直接复制】", font_size=11, bold=True, space_before=6)
prompt1 = (
    "竖屏9:16比例，微距特写，极具真实生活感的木质桌面。"
    "一双自然质感的手正在快速撕开一个零食包装袋，"
    "瞬间挤出一个外表金黄、泛着诱人油光的大鸡腿。"
    "暖色调侧方自然光打在鸡腿表面，清晰展现鸡皮的颗粒感和紧致度，"
    "画面带有微小抖动的手持实拍感，iPhone原相机质感，"
    "电影级超高清，食物摄影，极度诱人，浅景深。"
)
add_para(doc, prompt1, font_size=11, color=(0, 100, 0))

add_para(doc, "【运镜】极快推镜头，从撕包装直接怼到金黄鸡腿特写", font_size=10, color=(100, 100, 100))
add_para(doc, '【花字】画面正中弹出："饿了馋了？来个大鸡腿！"（明黄色大字，粗黑边）', font_size=10, color=(100, 100, 100))

add_divider(doc)

# 镜头2
add_para(doc, "镜头2：超强食欲的真实试吃（00:03 - 00:08）", font_size=13, bold=True, color=(50, 50, 50), space_before=10)

add_para(doc, "【口播文案】（大口咀嚼，含糊且享受）", font_size=11, bold=True)
add_para(doc, '"哇！这金黄的色泽，皮Q肉紧，咸香入骨，太解馋了！"', font_size=12, color=(200, 50, 0))

add_para(doc, "【即梦AI提示词 — 直接复制】", font_size=11, bold=True, space_before=6)
prompt2 = (
    "竖屏9:16比例，近景人像，一个穿着纯色卫衣的年轻女生在温馨居家环境中，"
    "手持金黄色的盐焗鸡腿，大口咬下，用力撕扯一大块鸡肉。"
    "表情极度享受，微微闭眼，咀嚼动作夸张真实，嘴角泛着油光。"
    "面部自然光影，真实皮肤质感，iPhone原相机实拍风格，"
    "居家暖色调背景，手持固定镜头感，让人看了就咽口水的食欲感，"
    "美食吃播风格，超高清。"
)
add_para(doc, prompt2, font_size=11, color=(0, 100, 0))

add_para(doc, "【运镜】固定手持感镜头，记录咬下鸡肉瞬间，头部随撕扯自然后仰", font_size=10, color=(100, 100, 100))

add_divider(doc)

# 镜头3
add_para(doc, "镜头3：核心卖点与肉质特写（00:08 - 00:12）", font_size=13, bold=True, color=(50, 50, 50), space_before=10)

add_para(doc, "【口播文案】（惊叹，语速加快）", font_size=11, bold=True)
add_para(doc, '"22年专注的纯正盐焗原味，你看这肉丝，连骨头里都是香的！"', font_size=12, color=(200, 50, 0))

add_para(doc, "【即梦AI提示词 — 直接复制】", font_size=11, bold=True, space_before=6)
prompt3 = (
    "竖屏9:16比例，极限微距特写，背景是虚化的粗盐粒。"
    "双手将鸡腿肉沿着纹理撕开，展示内部一丝一丝紧实白嫩的鸡肉，"
    "肉丝间渗出晶莹的汁水。高锐度，高细节，"
    "暖光精准勾勒肉质纤维的立体感，画面极其逼真，"
    "食物摄影风格，微距镜头，浅景深，超高清，"
    "仿佛能闻到盐焗的咸香味。"
)
add_para(doc, prompt3, font_size=11, color=(0, 100, 0))

add_para(doc, "【运镜】缓慢稳定的微距跟镜，聚焦肉质纤维细节", font_size=10, color=(100, 100, 100))

add_divider(doc)

# 镜头4（补充结尾）
add_para(doc, "镜头4：引导下单收尾（00:12 - 00:18）", font_size=13, bold=True, color=(50, 50, 50), space_before=10)

add_para(doc, "【口播文案】（语速快，紧迫感）", font_size=11, bold=True)
add_para(doc, '"这么大一个才几块钱，点下方链接直接拍，手慢就没了！"', font_size=12, color=(200, 50, 0))

add_para(doc, "【即梦AI提示词 — 直接复制】", font_size=11, bold=True, space_before=6)
prompt4 = (
    "竖屏9:16比例，中景，年轻女生穿纯色卫衣坐在居家环境中，"
    "手里举着咬了一半的金黄鸡腿，对着镜头自信微笑，"
    "另一只手指向下方（引导点击），表情满足又推荐的神态。"
    "背景温馨自然，暖色调，iPhone原相机实拍风格，"
    "真实生活感，美食吃播收尾，超高清。"
)
add_para(doc, prompt4, font_size=11, color=(0, 100, 0))

add_para(doc, '【花字】底部弹出："点击下方链接，立即下单"（红色+箭头动画）', font_size=10, color=(100, 100, 100))

add_divider(doc)

# ===== 剪映拼接指南 =====
add_para(doc, "四、剪映拼接指南", font_size=15, bold=True, color=(200, 50, 0), space_before=8)

steps = [
    "1. 在即梦分别生成 4 个片段（每个约 5 秒），下载到手机",
    "2. 打开剪映 → 新建项目 → 导入 4 个视频片段",
    "3. 按镜头1-4顺序排列，调整每段时长：3s + 5s + 4s + 6s",
    "4. 添加口播配音：即梦「配音生成」或剪映「文字朗读」，选少女音/活力音色",
    "5. 添加音效：剪映音效库搜索「撕包装」「咀嚼」「吞咽」",
    "6. BGM：剪映曲库搜索「卡点」「美食」「节奏感」，选动感打击乐",
    "7. 添加花字：使用剪映「文字模板」，选美食/种草类动态字幕",
    "8. 导出前用剪映「违禁词检测」功能扫描文案",
    "9. 导出设置：1080P，60fps，9:16竖屏",
]
for s in steps:
    add_para(doc, s, font_size=11, space_after=2)

add_divider(doc)

# ===== 发布建议 =====
add_para(doc, "五、抖音发布建议", font_size=15, bold=True, color=(200, 50, 0), space_before=8)

publish = [
    "发布时间：11:30-13:00 或 17:30-19:00（饭点流量高峰）",
    "标题示例：饿了馋了来个无穷盐焗大鸡腿！皮Q肉紧咸香入骨 #美食推荐 #零食测评",
    "话题标签：#无穷鸡腿 #盐焗鸡腿 #零食推荐 #吃货日常 #美食测评 #解馋零食",
    "挂车：绑定抖音小店或精选联盟商品链接",
    "互动引导：评论区置顶 '你们喜欢盐焗味还是蜜汁味？评论区告诉我'",
    "注意：视频左上角标注 '广告' 或 '含推广内容'（带货合规要求）",
]
for p_text in publish:
    add_para(doc, "· " + p_text, font_size=11, space_after=3)

output_path = "/mnt/d/Claude Code Data/无穷盐焗鸡腿_抖音带货视频脚本.docx"
doc.save(output_path)
print(f"文档已保存: {output_path}")
