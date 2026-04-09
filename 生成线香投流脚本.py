# -*- coding: utf-8 -*-
# 青澜香馆·沉萦系列  线香信息流投放视频脚本生成器
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def make_run(p, text, font_size=12, bold=False, color=None):
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_para(doc, text='', font_size=12, bold=False, align=None,
             color=None, space_before=2, space_after=4, indent=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    make_run(p, text, font_size, bold, color)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)
    return p


def add_field(doc, label, body, indent=0.5):
    """带标签的字段行，label加粗棕色，body正文色"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(22)
    make_run(p, label, 11, bold=True, color=(102, 51, 0))
    make_run(p, body, 11, color=(50, 50, 50))


def add_divider(doc):
    add_para(doc, '─' * 45, font_size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=(180, 180, 180))


# ======================================================
# 文档初始化
# ======================================================
doc = Document()
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin    = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

# ======================================================
# 大标题
# ======================================================
add_para(doc, '青澜香馆 · 沉萦系列  线香信息流投放视频脚本',
         font_size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(102, 51, 0))
add_para(doc, '投流专用  |  抖音合规版  |  视频时长：60秒  |  风格：新中式 / 禅意 / 质感',
         font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
         color=(150, 120, 80), space_before=0, space_after=6)

add_divider(doc)

# ======================================================
# 一、抖音投流合规要点
# ======================================================
add_para(doc, '一、抖音投流合规要点',
         font_size=15, bold=True, color=(102, 51, 0), space_before=8)
rules = [
    '✅  未使用极限词（最好、第一、唯一、顶级、100% 等）',
    '✅  未使用功效断言（助眠、安神、杀菌、净化空气、辟邪、转运等）',
    '✅  未贬低竞品，未使用「化学香危害」等恐吓营销话术',
    '✅  氛围描述采用主观情绪表达（如「陪你」「画个句号」）而非功效承诺',
    '✅  行动引导使用「喜欢可以拍」代替虚假催单话术',
    '⚠️  「天然原料」「植物配方」等说法需有检测报告支撑方可使用',
    '⚠️  「二十八道工序」等具体数字需与实际生产工艺相符',
    '⚠️  投流素材须标注「广告」标识，避免出现明火靠近易燃物的画面',
]
for r in rules:
    add_para(doc, r, font_size=11, space_after=3, indent=0.3)

add_divider(doc)

# ======================================================
# 二、全局拍摄设定
# ======================================================
add_para(doc, '二、全局拍摄设定',
         font_size=15, bold=True, color=(102, 51, 0), space_before=8)
for k, v in [
    ('画面比例', '9:16 竖屏短视频'),
    ('视觉风格', '新中式美学场景，自然侧光或逆光，放大线香烟雾丝缕质感与香材天然纹理，保持真实生活感，拒绝过度美化'),
    ('色彩调性', '偏暖黄低饱和，滤镜选胶片或暖阳，整体温润克制'),
    ('剪辑节奏', '钩子段落快切每镜2秒内，工艺与氛围段落慢节奏每镜3-5秒'),
    ('BGM   ', '禅意古风轻音乐 / 中式空灵，音量20%-30%'),
    ('字幕样式', '宋体或楷体，白色字幕黑色描边，底部居中'),
    ('转场方式', '全程叠化转场，0.5秒'),
]:
    add_field(doc, k + '：', v)

add_divider(doc)

# ======================================================
# 三、分镜脚本
# ======================================================
add_para(doc, '三、分镜脚本',
         font_size=15, bold=True, color=(102, 51, 0), space_before=8)

scenes = [
    {
        'title': '分镜1  钩子开篇（0-5秒）',
        'shot':  '极致特写 / 逆光丁