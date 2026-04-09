from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_text(cell, text, font_size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.bold = bold

def add_title(doc, text, font_size=18, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_info_line(doc, text, font_size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(2)
    return p

def add_section_title(doc, text, font_size=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_question(doc, text, font_size=12, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(24)
    return p

def add_blank_line(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">' \
              '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
              '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
              '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
              '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
              '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
              '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
              '</w:tblBorders>'
    from lxml import etree
    tblPr.append(etree.fromstring(borders))

# ==================== 试卷 ====================
doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 标题
add_title(doc, "小学一年级下学期数学测试卷", font_size=20)
add_info_line(doc, "满分：100分　　时间：60分钟", font_size=11)

# 姓名班级行
p = doc.add_paragraph()
run = p.add_run("姓名：______________　　班级：______________　　得分：______________")
run.font.size = Pt(11)
run.font.name = "宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.space_after = Pt(6)

# ===== 一、填空题 =====
add_section_title(doc, "一、填空题（每空 2 分，共 20 分）")

fill_questions = [
    "1. 从右边起，第一位是（　　）位，第二位是（　　）位，第三位是（　　）位。",
    "2. 56 是由（　　）个十和（　　）个一组成的。",
    "3. 和 79 相邻的两个数是（　　）和（　　）。",
    "4. 最大的两位数是（　　），最小的两位数是（　　）。",
    "5. 1 元 =（　　）角，　1 角 =（　　）分。",
    "6. 钟面上时针指向 8，分针指向 12，是（　　）时。",
]
for q in fill_questions:
    add_question(doc, q)

# ===== 二、口算题 =====
add_section_title(doc, "二、口算题（每题 1 分，共 20 分）")

calc_data = [
    ["36 + 5 =", "72 - 8 =", "45 + 30 =", "90 - 50 ="],
    ["28 + 7 =", "63 - 6 =", "54 + 20 =", "80 - 30 ="],
    ["15 + 9 =", "41 - 5 =", "37 + 40 =", "70 - 20 ="],
    ["46 + 8 =", "82 - 9 =", "25 + 60 =", "100 - 40 ="],
    ["53 + 9 =", "34 - 7 =", "68 + 10 =", "60 - 6 ="],
]

table = doc.add_table(rows=5, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table)
for i, row_data in enumerate(calc_data):
    for j, val in enumerate(row_data):
        set_cell_text(table.cell(i, j), val, font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        table.cell(i, j).paragraphs[0].paragraph_format.space_before = Pt(4)
        table.cell(i, j).paragraphs[0].paragraph_format.space_after = Pt(4)

# ===== 三、比较大小 =====
add_section_title(doc, "三、比较大小（填 >、< 或 =）（每题 2 分，共 10 分）")

compare_questions = [
    "1.  56  ○  65",
    "2.  78  ○  72 + 6",
    "3.  100  ○  99",
    "4.  43 - 8  ○  30",
    "5.  5 角  ○  48 分",
]
for q in compare_questions:
    add_question(doc, q)

# ===== 四、选择题 =====
add_section_title(doc, "四、选择题（每题 2 分，共 10 分）")

choice_questions = [
    ("1. 下面哪个数最接近 50？（　　）", "A. 38　　　B. 49　　　C. 55"),
    ("2. 小明有 35 颗糖，又买了 8 颗，一共有（　　）颗。", "A. 27　　　B. 43　　　C. 42"),
    ("3. 一个数个位上是 7，十位上是 3，这个数是（　　）。", "A. 73　　　B. 37　　　C. 307"),
    ("4. 用 2、0、5 组成的最大两位数是（　　）。", "A. 52　　　B. 50　　　C. 25"),
    ("5. 下列图形中（　　）不是长方形。", "A. 正方形　　B. 三角形　　C. 以上都不是"),
]
for stem, options in choice_questions:
    add_question(doc, stem)
    add_question(doc, options, indent=1.2)

# ===== 五、列竖式计算 =====
add_section_title(doc, "五、列竖式计算（每题 3 分，共 12 分）")

p = doc.add_paragraph()
run = p.add_run("1.  56 + 37 =　　　　　2.  82 - 45 =　　　　　3.  34 + 28 =　　　　　4.  71 - 36 =")
run.font.size = Pt(12)
run.font.name = "宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_after = Pt(4)

add_blank_line(doc, 5)

# ===== 六、画一画 =====
add_section_title(doc, "六、画一画（共 4 分）")
add_question(doc, "请在下面的空白处分别画出 1 个长方形和 1 个正方形。")
add_blank_line(doc, 4)

# ===== 七、解决问题 =====
add_section_title(doc, "七、解决问题（每题 4 分，共 24 分）")

word_problems = [
    "1. 学校合唱队有 45 人，又加入了 18 人，合唱队现在一共有多少人？",
    "2. 停车场原来有 72 辆车，开走了 25 辆，还剩多少辆？",
    "3. 小红有 36 张画片，小明比小红多 9 张，小明有多少张画片？",
    "4. 妈妈买了一个书包用了 45 元，买了一盒彩笔用了 18 元，一共用了多少元？",
    "5. 一班有 42 人，二班有 38 人，一班比二班多多少人？",
    "6. 小华看一本书，已经看了 28 页，还剩 34 页没看，这本书一共有多少页？",
]

for wp in word_problems:
    add_question(doc, wp)
    add_question(doc, "列式：______________________________", indent=1.0)
    add_question(doc, "答：________________________________", indent=1.0)
    add_blank_line(doc, 1)

output_path_exam = "/mnt/d/Claude Code Data/一年级下学期数学测试卷.docx"
doc.save(output_path_exam)
print(f"试卷已保存: {output_path_exam}")


# ==================== 答案 ====================
doc2 = Document()

section2 = doc2.sections[0]
section2.page_width = Cm(21)
section2.page_height = Cm(29.7)
section2.top_margin = Cm(2)
section2.bottom_margin = Cm(2)
section2.left_margin = Cm(2.5)
section2.right_margin = Cm(2.5)

add_title(doc2, "小学一年级下学期数学测试卷 — 参考答案", font_size=18)

# 一、填空题答案
add_section_title(doc2, "一、填空题")
answers_fill = [
    "1. 个位、十位、百位",
    "2. 5 个十和 6 个一",
    "3. 78 和 80",
    "4. 99、10",
    "5. 10 角、10 分",
    "6. 8 时",
]
for a in answers_fill:
    add_question(doc2, a)

# 二、口算题答案
add_section_title(doc2, "二、口算题")
calc_answers = [
    ["36 + 5 = 41", "72 - 8 = 64", "45 + 30 = 75", "90 - 50 = 40"],
    ["28 + 7 = 35", "63 - 6 = 57", "54 + 20 = 74", "80 - 30 = 50"],
    ["15 + 9 = 24", "41 - 5 = 36", "37 + 40 = 77", "70 - 20 = 50"],
    ["46 + 8 = 54", "82 - 9 = 73", "25 + 60 = 85", "100 - 40 = 60"],
    ["53 + 9 = 62", "34 - 7 = 27", "68 + 10 = 78", "60 - 6 = 54"],
]
table2 = doc2.add_table(rows=5, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table2)
for i, row_data in enumerate(calc_answers):
    for j, val in enumerate(row_data):
        set_cell_text(table2.cell(i, j), val, font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        table2.cell(i, j).paragraphs[0].paragraph_format.space_before = Pt(4)
        table2.cell(i, j).paragraphs[0].paragraph_format.space_after = Pt(4)

# 三、比较大小答案
add_section_title(doc2, "三、比较大小")
compare_answers = [
    "1.  56  <  65",
    "2.  78  =  72 + 6（72 + 6 = 78）",
    "3.  100  >  99",
    "4.  43 - 8  >  30（43 - 8 = 35）",
    "5.  5 角  >  48 分（5 角 = 50 分）",
]
for a in compare_answers:
    add_question(doc2, a)

# 四、选择题答案
add_section_title(doc2, "四、选择题")
choice_answers = [
    "1.  B（49 最接近 50）",
    "2.  B（35 + 8 = 43）",
    "3.  B（37）",
    "4.  A（52）",
    "5.  B（三角形不是长方形）",
]
for a in choice_answers:
    add_question(doc2, a)

# 五、竖式计算答案
add_section_title(doc2, "五、列竖式计算")
vertical_answers = [
    "1.  56 + 37 = 93",
    "2.  82 - 45 = 37",
    "3.  34 + 28 = 62",
    "4.  71 - 36 = 35",
]
for a in vertical_answers:
    add_question(doc2, a)

# 六、画一画答案
add_section_title(doc2, "六、画一画")
add_question(doc2, "长方形和正方形各画 1 个即可得分，注意正方形四边要等长。")

# 七、解决问题答案
add_section_title(doc2, "七、解决问题")
wp_answers = [
    ("1. 学校合唱队", "45 + 18 = 63（人）", "答：合唱队现在一共有 63 人。"),
    ("2. 停车场", "72 - 25 = 47（辆）", "答：还剩 47 辆车。"),
    ("3. 画片", "36 + 9 = 45（张）", "答：小明有 45 张画片。"),
    ("4. 买东西", "45 + 18 = 63（元）", "答：一共用了 63 元。"),
    ("5. 班级人数", "42 - 38 = 4（人）", "答：一班比二班多 4 人。"),
    ("6. 看书", "28 + 34 = 62（页）", "答：这本书一共有 62 页。"),
]
for title, formula, answer in wp_answers:
    add_question(doc2, title)
    add_question(doc2, "列式：" + formula, indent=1.0)
    add_question(doc2, answer, indent=1.0)

# 评分标准
add_section_title(doc2, "评分标准")
scoring = [
    "竖式计算：列式正确但答案算错扣 1 分",
    "应用题：列式正确但计算错误扣 1 分，没写'答'扣 1 分",
]
for s in scoring:
    add_question(doc2, "· " + s)

output_path_answer = "/mnt/d/Claude Code Data/一年级下学期数学测试卷_参考答案.docx"
doc2.save(output_path_answer)
print(f"答案已保存: {output_path_answer}")
